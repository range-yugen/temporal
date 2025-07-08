from temporalio import activity
import sqlite3
from datetime import datetime, timedelta
import random
from docx import Document
from docx2pdf import convert
import os
import uuid

OUTPUT_DIR = "../static/prescriptions"
BASE_URL = "http://localhost:8000"

os.makedirs(OUTPUT_DIR, exist_ok=True)

DB_PATH = "../clinic.db"

def get_connection():
    return sqlite3.connect(DB_PATH)

@activity.defn
async def check_doctor_availability(doctor_name: str) -> dict:
    conn = get_connection()
    cur = conn.cursor()
    now = datetime.now()
    weekday = now.strftime("%A")
    time_now = now.strftime("%H:%M:%S")

    cur.execute("""
        SELECT doctor_id FROM doctor_schedule
        WHERE LOWER(name) = LOWER(?)
        AND day_of_week = ?
        AND start_time <= ?
        AND end_time >= ?
    """, (doctor_name, weekday, time_now, time_now))

    row = cur.fetchone()
    conn.close()

    if row:
        return {
            "available": True,
            "doctor_id": row[0]
        }
    else:
        return {
            "available": False,
            "doctor_id": None
        }

@activity.defn
async def get_patient_by_phone(phone_number: str) -> dict:
    """Get patient information by phone number"""
    conn = get_connection()
    cur = conn.cursor()
    
    cur.execute("""
        SELECT patient_id, name, phone, gender, age, address FROM patients
        WHERE phone = ?
    """, (phone_number,))
    
    patient = cur.fetchone()
    conn.close()
    
    if patient:
        return {
            "patient_id": patient[0],
            "name": patient[1],
            "phone_number": patient[2],
            "gender": patient[3],
            "age": patient[4],
            "address": patient[5]
        }
    return None

@activity.defn
async def confirm_patient_appointment(patient_id: int, doctor_id: int) -> bool:
    conn = get_connection()
    cur = conn.cursor()

    # Get today's start and end time
    now = datetime.now()
    start_of_day = now.replace(hour=0, minute=0, second=0, microsecond=0)
    end_of_day = start_of_day + timedelta(days=1)

    cur.execute("""
        SELECT * FROM appointments
        WHERE patient_id = ? AND doctor_id = ?
        AND appointment_datetime >= ? AND appointment_datetime < ?
        AND status = 'scheduled'
        ORDER BY appointment_datetime ASC
        LIMIT 1
    """, (patient_id, doctor_id, start_of_day.isoformat(), end_of_day.isoformat()))

    appointment = cur.fetchone()
    conn.close()
    return appointment is not None

@activity.defn
async def register_patient(name: str, phone: str, gender: str, age: str, address: str) -> str:
    """
    Register a new patient with name, phone, gender, age and address.
    """
    conn = get_connection()
    cur = conn.cursor()

    # Check if patient already exists
    cur.execute("SELECT patient_id FROM patients WHERE phone = ?", (phone,))
    existing = cur.fetchone()
    if existing:
        conn.close()
        return f"Patient with phone {phone} already registered (patient_id: {existing[0]})."

    cur.execute("SELECT MAX(patient_id) FROM patients")
    max_id_row = cur.fetchone()
    max_id = max_id_row[0] if max_id_row[0] is not None else 0
    patient_id = max_id + 1

    # Insert new patient
    cur.execute("""
        INSERT INTO patients (patient_id, name, phone, gender, age, address)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (patient_id, name, phone, gender, age, address))

    conn.commit()
    conn.close()

    return f"Patient registered successfully with patient_id: {patient_id}"

@activity.defn
async def estimate_wait_time_for_walkin(doctor_id: int) -> int:
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT COUNT(*) FROM doctor_queue
        WHERE doctor_id = ? AND (seen IS NULL OR seen = 'no')
    """, (doctor_id,))
    
    queue_count = cur.fetchone()[0]
    conn.close()

    return queue_count * 15

@activity.defn
async def book_later_appointment(patient_id: int, doctor_id: int):
    conn = get_connection()
    cur = conn.cursor()
    now = datetime.now()
    
    # Step 1: Check if patient already has a future appointment with this doctor
    cur.execute("""
        SELECT 1 FROM appointments
        WHERE patient_id = ? AND doctor_id = ?
        AND appointment_datetime >= ? AND status = 'scheduled'
        LIMIT 1
    """, (patient_id, doctor_id, now.isoformat()))
    
    already_booked = cur.fetchone()
    if already_booked:
        conn.close()
        return "Patient already has a scheduled appointment with this doctor."

    # Step 2: Proceed to book next available slot
    weekday_today = now.strftime("%A")
    weekday_tomorrow = (now + timedelta(days=1)).strftime("%A")
    hour_now = now.hour

    is_morning = hour_now < 12

    # Decide schedule day and appointment date
    if is_morning:
        schedule_day = weekday_today
        appointment_date = now.date()
    else:
        schedule_day = weekday_tomorrow
        appointment_date = (now + timedelta(days=1)).date()

    # Fetch doctor's schedule
    cur.execute("""
        SELECT start_time, end_time FROM doctor_schedule
        WHERE doctor_id = ? AND day_of_week = ?
    """, (doctor_id, schedule_day))
    schedules = cur.fetchall()

    if not schedules:
        conn.close()
        return f"No schedule found for this doctor on {schedule_day} day."

    # Try to find an available 15-minute slot
    for start_str, end_str in schedules:
        start_dt = datetime.strptime(f"{appointment_date} {start_str}", "%Y-%m-%d %H:%M:%S")
        end_dt = datetime.strptime(f"{appointment_date} {end_str}", "%Y-%m-%d %H:%M:%S")

        current_slot = start_dt
        while current_slot + timedelta(minutes=15) <= end_dt:
            cur.execute("""
                SELECT 1 FROM appointments
                WHERE doctor_id = ? AND appointment_datetime = ? AND status = 'scheduled'
            """, (doctor_id, current_slot.isoformat()))
            conflict = cur.fetchone()

            if not conflict:
                # Found available slot
                cur.execute("""
                    INSERT INTO appointments (patient_id, doctor_id, appointment_datetime, status)
                    VALUES (?, ?, ?, 'scheduled')
                """, (patient_id, doctor_id, current_slot.isoformat()))
                conn.commit()
                conn.close()
                return (patient_id, doctor_id, str(current_slot))

            current_slot += timedelta(minutes=15)

    conn.close()
    return "All 15-minute slots are already booked for this doctor."

@activity.defn
async def add_to_walkin_queue(patient_id: int, doctor_id: int):
    """Add patient to walk-in queue if not already in the queue with seen = 'no'"""
    conn = get_connection()
    cur = conn.cursor()

    # Check if patient already in queue and not seen yet
    cur.execute("""
        SELECT 1 FROM doctor_queue
        WHERE patient_id = ? AND doctor_id = ? AND (seen IS NULL OR seen = 'no')
        LIMIT 1
    """, (patient_id, doctor_id))
    
    already_in_queue = cur.fetchone()
    
    if already_in_queue:
        conn.close()
        return False ## already available

    # Insert into queue
    cur.execute("""
        INSERT INTO doctor_queue (patient_id, doctor_id, queued_at, seen)
        VALUES (?, ?, ?, 'no')
    """, (patient_id, doctor_id, datetime.now().isoformat()))
    
    conn.commit()
    conn.close()
    return True ## added

@activity.defn
async def generate_prescription_slip(data: dict) -> dict:
    """
    Generate prescription slip from data dictionary
    """
    try:
        data["date"] = datetime.today().strftime('%Y-%m-%d')

        unique_id = str(uuid.uuid4())
        docx_path = os.path.join(OUTPUT_DIR, f"{unique_id}.docx")
        pdf_path = os.path.join(OUTPUT_DIR, f"{unique_id}.pdf")

        template_path = "../prescription/prescription_template.docx"
        
        # Check if template exists
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Prescription template not found at {template_path}")
        
        doc = Document(template_path)

        # Replace placeholders
        for p in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))
        
        doc.save(docx_path)
        convert(docx_path, pdf_path)

        return {
            "unique_id": unique_id,
            "pdf_url": f"{BASE_URL}/static/prescriptions/{unique_id}.pdf"
        }
        
    except KeyError as e:
        raise ValueError(f"Missing required patient data field: {e}")
    except Exception as e:
        raise Exception(f"Error generating prescription slip: {str(e)}")


@activity.defn
async def prescription_with_diagnosis(unique_id: str, diagnosis: str, medicines: list) -> str:
    # Original draft paths
    docx_path = os.path.join(OUTPUT_DIR, f"{unique_id}.docx")
    
    # Generate new unique ID for final prescription
    final_unique_id = f"{unique_id}_final"
    final_docx_path = os.path.join(OUTPUT_DIR, f"{final_unique_id}.docx")
    final_pdf_path = os.path.join(OUTPUT_DIR, f"{final_unique_id}.pdf")

    if not os.path.exists(docx_path):
        raise FileNotFoundError("Draft prescription not found")

    # Load the original document
    doc = Document(docx_path)

    # Locate the 'Rx -' paragraph
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() in {"rx -", "rx:", "rx"}:
            insert_index = i + 1
            break
    else:
        insert_index = len(doc.paragraphs)

    # Clear existing content after Rx and rebuild properly
    # Remove paragraphs after Rx
    paragraphs_to_remove = doc.paragraphs[insert_index:]
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    # Add diagnosis first
    if diagnosis:
        doc.add_paragraph(f"Diagnosis: {diagnosis}")
        doc.add_paragraph("")  # Empty line for spacing

    # Add medicines section
    if medicines:
        doc.add_paragraph("Medicines:")
        for medicine in medicines:
            doc.add_paragraph(f"- {medicine}")

    # Save as final prescription with new unique ID
    doc.save(final_docx_path)
    convert(final_docx_path, final_pdf_path)

    return f"{BASE_URL}/static/prescriptions/{final_unique_id}.pdf"


@activity.defn
async def get_random_diagnosis_and_medicines() -> dict:
    """
    Get a random diagnosis and its associated medicines from the database
    Returns: dict with 'diagnosis' and 'medicines' (as list)
    """
    conn = get_connection()
    cur = conn.cursor()
    
    try:
        # Get all diagnosis and medicines from the table
        cur.execute("SELECT diagnosis, medicines FROM diagnosis_medicines")
        all_records = cur.fetchall()
        
        if not all_records:
            # Fallback if no records found
            return {
                "diagnosis": "General Consultation",
                "medicines": ["Multivitamin", "Adequate Rest"]
            }
        
        # Select a random record
        random_record = random.choice(all_records)
        diagnosis = random_record[0]
        medicines_str = random_record[1]
        
        # Convert comma-separated medicines to list
        medicines = [medicine.strip() for medicine in medicines_str.split(',')]
        
        return {
            "diagnosis": diagnosis,
            "medicines": medicines
        }
        
    except Exception as e:
        # Fallback in case of any error
        return {
            "diagnosis": "General Health Check",
            "medicines": ["As advised by doctor"]
        }
    
    finally:
        conn.close()